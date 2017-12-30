from bottle import Bottle, route, run, request
import requests
import urllib
from docx.shared import Inches
from bs4 import BeautifulSoup,NavigableString,Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
app = Bottle()

@app.hook('after_request')
def enable_cors():
    """
    You need to add some headers to each request.
    Don't use the wildcard '*' for Access-Control-Allow-Origin in production.
    """
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'PUT, GET, POST, DELETE, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Origin, Accept, Content-Type, X-Requested-With, X-CSRF-Token'


@app.route('/crawl', method=['OPTIONS', 'GET'])
def crawl():
	url = request.query.url
	source_code = requests.get(url).text
	soup = BeautifulSoup(source_code,"html.parser")
	#document = Document()
	for heading in soup.find('h1'):
		#(heading.string).encode('ascii', 'ignore')
		filename = heading.string
		print (heading.string)
		#filename = heading.string.replace('/',"")
	#story = open("/home/akshay/Desktop/Story/" + filename + ".docx","rb")	
	document = Document()
	document.add_heading(filename,0)
	#	document.add_paragraph(filename)
	
	i=1
	#j=0
	for content in soup.select('p'):
		
	

		for pic in content.select('img'):
			#urlpic = pic.get('src')
			#a = urlpic.replace('/',"")
			#urllib.request.urlretrieve(urlpic,str(i)+".jpg")
			#urllib.urlretrieve(urlpic,str(i)+".jpg")
			#document.add_picture(str(i)+".jpg")
			#i = i + 1;
			try:
				urlpic = pic.get('src')
				urllib.request.urlretrieve(urlpic,str(i)+".jpg")
				document.add_picture(str(i)+".jpg")
				i = i + 1;

			except ValueError :
				
				urlpic = pic.get('src')
				startIndexCount = url.find("/") + 2
				endIndex = url.find("/", startIndexCount)
				mainWebsiteUrl = url[:endIndex]

				#print(mainWebsiteUrl)

	            
				actualUrl = mainWebsiteUrl + urlpic
				#print(actualUrl)
				urllib.request.urlretrieve(actualUrl,str(i)+".jpg")
				document.add_picture(str(i)+".jpg")
				i = i + 1;

		asd = content.find_next_sibling()
		if asd is not None :
			#print (asd)
			if asd.name == 'div' :
				
				for pic in asd.select('img'):
					#print (pic)
					
					#print(url)
					#url = "http://www.hindustantimes.com" + url 
					#a = url.replace('/',"")
					#urllib.request.urlretrieve(url,a)
					#urllib.urlretrieve(url,str(i)+".jpg")
					try:
						urlpic = pic.get('src')
						urllib.request.urlretrieve(urlpic,str(i)+".jpg")
						document.add_picture(str(i)+".jpg")
						i = i + 1
					except ValueError :
				
						urlpic = pic.get('src')
						startIndexCount = url.find("/") + 2
						endIndex = url.find("/", startIndexCount)
						mainWebsiteUrl = url[:endIndex]

						#print(mainWebsiteUrl)

	            
						actualUrl = mainWebsiteUrl + urlpic
						#print(actualUrl)
						urllib.request.urlretrieve(actualUrl,str(i)+".jpg")
						document.add_picture(str(i)+".jpg")
						i = i + 1
				
			elif asd.name == 'pre' :

				paragraph=document.add_paragraph(content.text)
				paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
				document.add_paragraph(asd.text)
				
				continue
			elif asd.name == 'ul' :
				paragraph=document.add_paragraph(content.text)
				paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
				paragraph.keep_together = True
				paragraph=document.add_paragraph(asd.text)
				paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
				paragraph.keep_together = True
				#paragraph=document.add_paragraph(asd.text, style='ListBullet') paragraph.alignment = 3
				continue

		
		paragraph=document.add_paragraph(content.text)
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
		paragraph.keep_together = True
		

			
		#print(content.find_next_sibling('div'))
		#story.write(content.text)
		#story.write("\n")
	
	
	document.save("/home/akshay/Desktop/Story/test/document.docx")
	return json.dumps("save file in path /home/akshay/Desktop/Story/test/document.docx")


if __name__ == "__main__":
	run(app, host='0.0.0.0', port=8080)
